import os
import docx
import zipfile
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


link = []
dictForm = {}
dictForm['name']=None
dictForm['email'] = None
dictForm['phone'] = None
dictForm['linkedin'] = None
dictForm['lineCount'] = 0
dictForm['textCount'] = 0
dictForm['tableCount'] = 0
dictForm['imgCount'] = 0

fontDetails = {}

#Read docx File
doc = docx.Document('Resume_1.docx')
paras = []

for block in iter_block_items(doc):
    if isinstance(block, Paragraph):
        paras.append(block.text.lstrip().rstrip().lstrip('\t').rstrip('\t').lstrip().rstrip())
    else:
        for row in block.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paras.append(paragraph.text)
text = '\n'.join(paras)
#End Reading 

for par in doc.paragraphs:
    #Find images count
    if 'graphicData' in par._p.xml:
        dictForm['imgCount']+=1
    #End image counting
    for run in par.runs:
        fsize = run.font.size
        fname = run.font.name
        if fname is not None and fsize is not None:
            fsize = fsize/12700
            if (fname in fontDetails) and (fsize not in fontDetails[fname]):
                fontDetails[fname] = fontDetails[fname].append(fsize)
            else:
                fontDetails[fname] = [fsize]
        
#Find table in docx
dictForm['tableCount'] = len(doc.tables)
#Find urls in Docx
rels = doc.part.rels
for rel in rels:
    if (rels[rel].reltype == RT.HYPERLINK):
        link.append(rels[rel]._target)
#End finding links

#Removing new-lines '\n' form the file to count only the words
flag = True

text1 = text
text2 = ""
while(flag):
    text2 = text1.replace("\n"," ").replace("  "," ")
    if(text2 == text1):
        flag = False
    else:
        text1 = text2

#End cleaning the "text" with '\n'

#Finding name, email, and linkedin profile
contentList = text.split("\n")

if(contentList[0].rstrip().lstrip() == "RESUME"):
    dictForm['name'] = contentList[2].lstrip().rstrip()
else:
    if text2.split()[0] == text2.split()[1]:
        dictForm['name'] = text2.split()[0].lstrip().rstrip()
    else:
        dictForm['name'] = text2.split()[0].lstrip().rstrip() + " " + text2.split()[1].lstrip().rstrip()


for i in link:
    if "linkedin" in i:
        dictForm['linkedin'] = i.lstrip().rstrip()
    if "@" in i:
        dictForm['email'] = i[str(i).find(":"):].lstrip().rstrip().strip(":")
#End find name,etc

#Find total number of Lines in the file
dictForm['lineCount'] = text.count("\n")
#End search for total Lines

#Finding the number of words"
dictForm['textCount'] = len(text2.strip())
#End

print(dictForm)
print(fontDetails)
